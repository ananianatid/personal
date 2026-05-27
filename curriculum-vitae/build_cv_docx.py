from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "cv-anatide-anani.docx"

INK = "111111"
NAVY = "182744"
BODY = "4A4A4A"
MUTED = "666666"
LINE = "DDDDDD"
SIDEBAR = "F7F7F6"
TAG_BG = "F2F5FA"
TAG_BORDER = "D8DEE8"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def font(run, size, color=BODY, bold=False, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.font.bold = bold
    return run


def run(paragraph, text, size=9, color=BODY, bold=False, name="Arial"):
    return font(paragraph.add_run(text), size, color, bold, name)


def pf(paragraph, before=0, after=3, line=1.1, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, top=None, right=None, bottom=None, left=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, value in {
        "top": top,
        "right": right,
        "bottom": bottom,
        "left": left,
    }.items():
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        if value is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), value.get("sz", "4"))
            element.set(qn("w:color"), value.get("color", LINE))


def no_borders(cell):
    borders(cell, top=None, right=None, bottom=None, left=None)


def margins(cell, top=80, right=80, bottom=80, left=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("end", right), ("bottom", bottom), ("start", left)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def para_rule(paragraph, color=LINE, size="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def section(cell, title, before=10):
    p = cell.add_paragraph()
    pf(p, before=before, after=7, line=1)
    run(p, "//  ", 8, NAVY, True, "Courier New")
    run(p, title.upper(), 7.4, NAVY, True, "Courier New")
    para_rule(p)


def profile(cell):
    p = cell.add_paragraph()
    pf(p, after=6, line=1.42)
    pieces = [
        ("Passionné d'informatique depuis le lycée, j'ai démarré avec ", False),
        ("Arduino et l'IoT", True),
        (" avant de plonger dans le développement web. Aujourd'hui en 3", False),
        ("ème", False),
        (" année de Licence en Génie Logiciel, je conçois des applications avec une approche ", False),
        ("architecture propre", True),
        (" et ", False),
        ("expérience utilisateur.", True),
    ]
    for text, bold in pieces:
        run(p, text, 8.8, INK if bold else BODY, bold)


def language(cell, name, level):
    row = cell.add_table(rows=1, cols=2)
    row.autofit = False
    row.columns[0].width = Cm(3.5)
    row.columns[1].width = Cm(1.8)
    for c in row.rows[0].cells:
        no_borders(c)
        margins(c, top=35, bottom=35, left=0, right=0)
        borders(c, bottom={"sz": "3", "color": LINE})
    p = row.rows[0].cells[0].paragraphs[0]
    pf(p, after=0, line=1)
    run(p, name, 8.3, INK)
    p = row.rows[0].cells[1].paragraphs[0]
    pf(p, after=0, line=1, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run(p, level, 7.2, NAVY, name="Courier New")


def tag_table(cell, tags, cols=2, size=6.4):
    rows = (len(tags) + cols - 1) // cols
    table = cell.add_table(rows=rows, cols=cols)
    table.autofit = False
    for row in table.rows:
        for c in row.cells:
            shade(c, WHITE)
            no_borders(c)
            margins(c, top=22, bottom=22, left=35, right=35)
    for idx, value in enumerate(tags):
        c = table.rows[idx // cols].cells[idx % cols]
        shade(c, TAG_BG)
        borders(
            c,
            top={"sz": "3", "color": TAG_BORDER},
            right={"sz": "3", "color": TAG_BORDER},
            bottom={"sz": "3", "color": TAG_BORDER},
            left={"sz": "3", "color": TAG_BORDER},
        )
        p = c.paragraphs[0]
        pf(p, after=0, line=1, align=WD_ALIGN_PARAGRAPH.CENTER)
        run(p, value, size, NAVY, name="Courier New")


def skill_group(cell, title, tags):
    p = cell.add_paragraph()
    pf(p, before=6, after=4, line=1)
    run(p, title.upper(), 6.8, NAVY, True, "Courier New")
    tag_table(cell, tags)


def simple_list(cell, values):
    for value in values:
        p = cell.add_paragraph()
        pf(p, after=3, line=1.1)
        run(p, value, 8.4, BODY)


def entry(cell, title, date=None, org=None, desc=None, tags=None, add_sep=True):
    table = cell.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.2)
    table.columns[1].width = Cm(3.4)
    for c in table.rows[0].cells:
        no_borders(c)
        margins(c, top=0, bottom=0, left=0, right=0)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = table.rows[0].cells[0].paragraphs[0]
    pf(p, after=0, line=1)
    run(p, title, 9.3, INK, True)
    if date:
        p = table.rows[0].cells[1].paragraphs[0]
        pf(p, after=0, line=1, align=WD_ALIGN_PARAGRAPH.RIGHT)
        run(p, date, 7.4, NAVY, name="Courier New")
    if org:
        p = cell.add_paragraph()
        pf(p, after=5, line=1)
        run(p, org, 7.3, MUTED, name="Courier New")
    if desc:
        p = cell.add_paragraph()
        pf(p, after=6, line=1.25)
        run(p, desc, 8.3, BODY)
    if tags:
        tag_table(cell, tags, cols=5, size=6.1)
    if add_sep:
        p = cell.add_paragraph()
        pf(p, before=6, after=8, line=1)
        para_rule(p)


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(0.18)
    sec.bottom_margin = Cm(0.18)
    sec.left_margin = Cm(0.18)
    sec.right_margin = Cm(0.18)
    sec.header_distance = Cm(0.1)
    sec.footer_distance = Cm(0.1)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(8.5)
    normal.font.color.rgb = rgb(BODY)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.1


def build():
    doc = Document()
    configure(doc)

    header = doc.add_table(rows=1, cols=1)
    header.autofit = False
    header.columns[0].width = Cm(20.64)
    h = header.cell(0, 0)
    shade(h, WHITE)
    borders(h, bottom={"sz": "5", "color": LINE})
    margins(h, top=500, bottom=420, left=120, right=120)

    p = h.paragraphs[0]
    pf(p, after=4, line=0.95, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Anatide ", 38, INK, True, "Arial Narrow")
    run(p, "ANANI", 38, NAVY, True, "Arial Narrow")

    p = h.add_paragraph()
    pf(p, after=8, line=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Développeur Web Full-Stack", 11.2, MUTED, True)

    p = h.add_paragraph()
    pf(p, after=16, line=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "DISPONIBLE POUR UN STAGE  ·  LOMÉ, TOGO", 7.3, NAVY, name="Courier New")

    p = h.add_paragraph()
    pf(p, after=0, line=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "ananianatid@gmail.com   ·   +228 70 07 68 29   ·   github.com/ananianatid", 8.2, NAVY)

    body = doc.add_table(rows=1, cols=2)
    body.autofit = False
    body.columns[0].width = Cm(6.85)
    body.columns[1].width = Cm(13.79)
    left, right = body.rows[0].cells
    for c in (left, right):
        no_borders(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade(left, SIDEBAR)
    shade(right, WHITE)
    borders(left, right={"sz": "5", "color": LINE})
    margins(left, top=520, bottom=420, left=820, right=520)
    margins(right, top=520, bottom=420, left=820, right=820)

    section(left, "Profil", before=0)
    profile(left)

    section(left, "Langues", before=16)
    language(left, "Français", "Avancé")
    language(left, "Anglais", "B2")
    language(left, "Ewe", "Natif")

    section(left, "Compétences", before=18)
    skill_group(left, "Backend", ["Laravel 10/12", "Statamic", "PHP 8+", "Node.js", "Express", "Sequelize", "Filament v3", "API REST", "Laravel Passport", "Sanctum"])
    skill_group(left, "Frontend & Mobile", ["JavaScript ES6+", "Expo", "HTML5", "CSS3", "Tailwind", "Bootstrap"])
    skill_group(left, "Bases de données", ["MySQL", "SQLite", "PostgreSQL"])
    skill_group(left, "Outils & DevOps", ["Git", "GitHub", "GitLab", "Vite", "Composer", "npm", "Supabase", "WordPress", "Docker"])
    skill_group(left, "Fondements & Scripting", ["Java", "Python", "C", "Bash"])

    section(left, "Loisirs", before=18)
    simple_list(left, ["Modélisation 3D (Blender)", "Échecs (1350 Elo)", "Lecture"])

    section(right, "Expérience", before=0)
    entry(
        right,
        "Gestionnaire de site WordPress",
        "2024 — Présent",
        "UFW Verein · ONG internationale",
        "Administration et maintenance du site institutionnel en production.",
    )

    section(right, "Formation", before=10)
    entry(
        right,
        "Licence en Génie Logiciel — L3",
        "2023 — En cours",
        "Université · Lomé, Togo",
        "Formation en génie logiciel — C, Python, Java, JavaScript, PHP. Projets académiques orientés conception et développement d'applications.",
    )
    entry(
        right,
        "Certification OPEM B",
        "2022 — 2023",
        "Conception d'une maison à éclairage connecté Bluetooth",
        "Arduino, AutoCAD, Blender, Fritzing et TinkerCAD.",
    )
    entry(
        right,
        "Certification OPEM A",
        "2021 — 2022",
        "Conception d'un portail à ouverture automatique",
        "Conception et programmation d'un système de portail automatisé avec Arduino.",
    )
    entry(right, "Baccalauréat Série D — Mention Bien", "2020 — 2023", "Lomé, Togo", add_sep=False)

    section(right, "Projets", before=14)
    entry(
        right,
        "logements_Mokpokpo",
        None,
        None,
        "Gestion des résidences universitaires — Filament v3, scoring de priorité pour l'attribution, contrats et paiements.",
        ["Laravel 12", "Filament v3", "PHP 8+", "MySQL", "Tailwind"],
    )
    entry(
        right,
        "collegeDirectory",
        None,
        None,
        "Annuaire universitaire avec double panel admin/étudiant. SSO OAuth2 / OpenID Connect via Laravel Passport et pattern shadow user.",
        ["Laravel", "Filament v3", "Passport", "OAuth2", "MySQL"],
    )
    entry(
        right,
        "App Atelier de Couture",
        None,
        None,
        "Application mobile offline-first pour commandes et mesures clients d'un atelier de couture. Modélisation dynamique des mesures.",
        ["React Native", "Expo", "SQLite", "Laravel", "MySQL"],
    )
    entry(
        right,
        "UFW-verein",
        None,
        None,
        "Administration et maintenance du site institutionnel en production. Contenu éditorial, modules, performances et incidents.",
        ["WordPress", "PHP"],
    )
    entry(
        right,
        "OQSF Bénin",
        None,
        None,
        "Portail institutionnel de l'Observatoire de la Qualité des Services Financiers : plaintes, ressources réglementaires et actualités.",
        ["Statamic", "Laravel", "PHP", "Tailwind"],
        add_sep=False,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
